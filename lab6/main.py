import tkinter as tk
from tkinter import messagebox

from room.room import Room
from apartment.apartment import Apartment
from building.building import Building

from docx import Document
from openpyxl import Workbook

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Многоэтажный дом")

        self.building = Building()
        self.current_apartment = Apartment()

        self.create_widgets()

    def create_widgets(self):
        # ввод
        tk.Label(text="Длина").grid(row=0, column=0)
        tk.Label(text="Ширина").grid(row=1, column=0)
        tk.Label(text="Высота").grid(row=2, column=0)

        self.length = tk.Entry()
        self.width = tk.Entry()
        self.height = tk.Entry()

        self.length.grid(row=0, column=1)
        self.width.grid(row=1, column=1)
        self.height.grid(row=2, column=1)

        # кнопки
        tk.Button(text="Добавить комнату", command=self.add_room).grid(row=3, column=0, columnspan=2)
        tk.Button(text="Новая квартира", command=self.new_apartment).grid(row=4, column=0, columnspan=2)
        tk.Button(text="Рассчитать дом", command=self.calculate).grid(row=5, column=0, columnspan=2)

        tk.Button(text="Сохранить DOCX", command=self.save_docx).grid(row=6, column=0)
        tk.Button(text="Сохранить XLSX", command=self.save_xlsx).grid(row=6, column=1)

        # список комнат
        tk.Label(text="Комнаты").grid(row=0, column=2)
        self.room_list = tk.Listbox(width=40, height=10)
        self.room_list.grid(row=1, column=2, rowspan=5)

        # список квартир
        tk.Label(text="Квартиры").grid(row=0, column=3)
        self.apartment_list = tk.Listbox(width=40, height=10)
        self.apartment_list.grid(row=1, column=3, rowspan=5)

        # результат
        self.result = tk.Label(text="")
        self.result.grid(row=7, column=0, columnspan=4)

    def add_room(self):
        try:
            l = float(self.length.get())
            w = float(self.width.get())
            h = float(self.height.get())

            if l <= 0 or w <= 0 or h <= 0:
                raise ValueError

            room = Room(l, w, h)
            self.current_apartment.add_room(room)

            self.room_list.insert(
                tk.END,
                f"{l}x{w}x{h} | {room.area():.1f} м²"
            )

        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числа")

    def new_apartment(self):
        if not self.current_apartment.rooms:
            messagebox.showwarning("Внимание", "Нет комнат")
            return

        self.building.add_apartment(self.current_apartment)

        index = len(self.building.apartments)
        area = self.current_apartment.total_area()

        self.apartment_list.insert(
            tk.END,
            f"Квартира {index}: {area:.1f} м²"
        )

        # новая квартира
        self.current_apartment = Apartment()
        self.room_list.delete(0, tk.END)

    def calculate(self):
        area = self.building.total_area()
        heat = self.building.total_heat()

        self.result.config(
            text=f"Дом: {area:.1f} м² | {heat:.1f} Вт"
        )

    def save_docx(self):
        doc = Document()
        doc.add_heading("Отчёт по дому", 0)

        for i, apt in enumerate(self.building.apartments, 1):
            doc.add_paragraph(f"Квартира {i}: {apt.total_area():.1f} м²")

        doc.add_paragraph(f"Общий метраж: {self.building.total_area():.1f} м²")
        doc.add_paragraph(f"Тепло: {self.building.total_heat():.1f} Вт")

        doc.save("building.docx")

    def save_xlsx(self):
        wb = Workbook()
        ws = wb.active

        ws["A1"] = "Квартира"
        ws["B1"] = "Площадь"

        for i, apt in enumerate(self.building.apartments, 1):
            ws[f"A{i+1}"] = f"Квартира {i}"
            ws[f"B{i+1}"] = apt.total_area()

        ws["A10"] = "Итого"
        ws["B10"] = self.building.total_area()

        wb.sav


("building.xlsx")

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()